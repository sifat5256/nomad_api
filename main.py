import random
import os
import base64

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse
from sendgrid import SendGridAPIClient
from sendgrid.helpers.mail import Mail, Attachment, FileContent, FileName, FileType, Disposition
from pydantic import BaseModel
from PyPDF2 import PdfMerger
import os
import uuid
import base64
from fastapi import UploadFile, File, FastAPI
from fastapi import BackgroundTasks
from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse
import os
import shutil
import uuid
import subprocess
from docx2pdf import convert

from fpdf import FPDF
from docx import Document

app = FastAPI()

SENDGRID_API_KEY = os.getenv("SENDGRID_API_KEY")
verification_codes = {}

# ---------------- Models ---------------- #

class EmailRequest(BaseModel):
    email: str

class VerificationRequest(BaseModel):
    email: str
    code: str

class TextEmailRequest(BaseModel):
    email: str
    subject: str
    message: str

# ---------------- Helpers ---------------- #

def generate_verification_code():
    return f"{random.randint(100000, 999999)}"

# ---------------- Routes ---------------- #

# ✅ 1. Merge PDF files
@app.post("/merge-pdf/")
async def merge_pdf(files: list[UploadFile] = File(...)):
    merger = PdfMerger()
    filenames = []

    for file in files:
        contents = await file.read()
        with open(file.filename, "wb") as f:
            f.write(contents)
        merger.append(file.filename)
        filenames.append(file.filename)

    output_filename = "merged.pdf"
    merger.write(output_filename)
    merger.close()

    for file in filenames:
        os.remove(file)

    return FileResponse(output_filename, media_type='application/pdf', filename="merged.pdf")

# ✅ 2. Send verification code via email
@app.post("/send-code/")
async def send_code(data: EmailRequest):
    email = data.email
    code = generate_verification_code()
    verification_codes[email] = code

    message = Mail(
        from_email=("no-reply@techapppartners.com", "Nomad"),
        to_emails=email,
        subject="Your Verification Code",
        html_content=f"<p>Your verification code is: <strong>{code}</strong></p>"
    )

    try:
        sg = SendGridAPIClient(SENDGRID_API_KEY)
        response = sg.send(message)
        if response.status_code not in range(200, 300):
            raise HTTPException(status_code=500, detail="Failed to send email")
        return {"message": "Verification code sent"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ✅ 3. Verify email code
@app.post("/verify-code/")
async def verify_code(data: VerificationRequest):
    email = data.email
    code = data.code
    stored_code = verification_codes.get(email)

    if stored_code is None:
        return {"verified": False, "message": "No code sent to this email"}
    if stored_code == code:
        del verification_codes[email]
        return {"verified": True, "message": "Email verified successfully"}
    else:
        return {"verified": False, "message": "Invalid code"}

# ✅ 4. Send text email
@app.post("/send-email/")
async def send_email(data: TextEmailRequest):
    message = Mail(
        from_email=("no-reply@techapppartners.com", "Nomad"),
        to_emails=data.email,
        subject=data.subject,
        html_content=f"<p>{data.message}</p>"
    )

    try:
        sg = SendGridAPIClient(SENDGRID_API_KEY)
        response = sg.send(message)
        if response.status_code not in range(200, 300):
            raise HTTPException(status_code=500, detail="Failed to send email")
        return {"message": "Email sent successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# ✅ 5. Send email with file attachments
@app.post("/send-email-with-files/")
async def send_email_with_files(
    email: str = Form(...),
    subject: str = Form(...),
    message: str = Form(...),
    files: list[UploadFile] = File(...)
):
    mail = Mail(
        from_email=("no-reply@techapppartners.com", "Nomad"),
        to_emails=email,
        subject=subject,
        html_content=f"<p>{message}</p>"
    )

    attachments = []
    for file in files:
        file_data = await file.read()
        encoded_file = base64.b64encode(file_data).decode()

        attachment = Attachment()
        attachment.file_content = FileContent(encoded_file)
        attachment.file_type = FileType(file.content_type)
        attachment.file_name = FileName(file.filename)
        attachment.disposition = Disposition("attachment")

        attachments.append(attachment)

    mail.attachment = attachments

    try:
        sg = SendGridAPIClient(SENDGRID_API_KEY)
        response = sg.send(mail)
        if response.status_code not in range(200, 300):
            raise HTTPException(status_code=500, detail="Failed to send email")
        return {"message": f"Email sent with {len(attachments)} attachment(s)"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
def docx_to_pdf(file_path, output_path):
    doc = Document(file_path)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for para in doc.paragraphs:
        pdf.multi_cell(0, 10, para.text)
    pdf.output(output_path)

def txt_to_pdf(file_path, output_path):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    with open(file_path, "r", encoding="utf-8") as f:
        for line in f:
            pdf.multi_cell(0, 10, line.strip())
    pdf.output(output_path)


def docx_to_pdf(input_path, output_path):
    doc = Document(input_path)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for para in doc.paragraphs:
        pdf.multi_cell(0, 10, para.text)
    pdf.output(output_path)



@app.post("/convert-doc-to-pdf/")
async def convert_doc_to_pdf(file: UploadFile = File(...), background_tasks: BackgroundTasks = None):
    if not file.filename.endswith((".docx", ".doc")):
        raise HTTPException(status_code=400, detail="File must be .docx or .doc")

    temp_dir = "temp_docs"
    os.makedirs(temp_dir, exist_ok=True)

    # Generate unique filename
    unique_id = str(uuid.uuid4())
    input_path = os.path.join(temp_dir, f"{unique_id}.docx")
    output_path = os.path.join(temp_dir, f"{unique_id}.pdf")

    # Save uploaded file
    with open(input_path, "wb") as f:
        contents = await file.read()
        f.write(contents)

    # Convert .docx to PDF
    try:
        docx_to_pdf(input_path, output_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {str(e)}")

    # Clean up temp files after sending response
    def cleanup():
        try:
            os.remove(input_path)
            os.remove(output_path)
        except:
            pass

    background_tasks.add_task(cleanup)

    return FileResponse(output_path, media_type="application/pdf", filename=os.path.basename(output_path))


    